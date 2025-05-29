from flask import Flask, render_template, request, jsonify, send_file, send_from_directory
import psycopg2
from flask_cors import CORS
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, ListFlowable, ListItem
from reportlab.lib.styles import getSampleStyleSheet
from docx import Document
from docx.shared import Inches
import io
import logging
import os
import telebot
from flask_socketio import SocketIO

app = Flask(__name__, static_folder='static')
CORS(app)

# Set up logging
logging.basicConfig(level=logging.DEBUG)
logger = logging.getLogger(__name__)

# Initialize SocketIO for WebSocket support
socketio = SocketIO(app, cors_allowed_origins='*')

# Initialize Telegram bot with environment variables
BOT_TOKEN = os.environ.get('BOT_TOKEN')
CHAT_ID = os.environ.get('CHAT_ID')
bot = telebot.TeleBot(BOT_TOKEN)

# Initialize Postgres database
def init_db():
    DATABASE_URL = os.environ['DATABASE_URL']
    with psycopg2.connect(DATABASE_URL, sslmode='require') as conn:
        cursor = conn.cursor()
        # Create table if it doesn't exist
        cursor.execute('''
            CREATE TABLE IF NOT EXISTS recipes (
                id SERIAL PRIMARY KEY,
                name TEXT NOT NULL,
                type TEXT NOT NULL,
                ingredients TEXT NOT NULL,
                instructions TEXT NOT NULL,
                remarks TEXT,
                image TEXT
            )
        ''')
        # Insert sample data (skip if already exists)
        cursor.execute('''
            INSERT INTO recipes (name, type, ingredients, instructions, remarks, image)
            VALUES (%s, %s, %s, %s, %s, %s)
            ON CONFLICT DO NOTHING
        ''', ('Test Main Meal', 'main_meal', 'Ingredient1,Ingredient2', 'Step 1. Step 2.', 'Sample remark', 'https://via.placeholder.com/150'))
        cursor.execute('''
            INSERT INTO recipes (name, type, ingredients, instructions, remarks, image)
            VALUES (%s, %s, %s, %s, %s, %s)
            ON CONFLICT DO NOTHING
        ''', ('Test Baking', 'baking', 'Flour,Sugar', 'Mix. Bake.', 'Baking note', 'https://via.placeholder.com/150'))
        cursor.execute('''
            INSERT INTO recipes (name, type, ingredients, instructions, remarks, image)
            VALUES (%s, %s, %s, %s, %s, %s)
            ON CONFLICT DO NOTHING
        ''', ('Test Desert', 'desert', 'Cream,Fruit', 'Chill. Serve.', 'Sweet note', 'https://via.placeholder.com/150'))
        conn.commit()

init_db()

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/main_meals')
def main_meals():
    DATABASE_URL = os.environ['DATABASE_URL']
    with psycopg2.connect(DATABASE_URL, sslmode='require') as conn:
        cursor = conn.cursor()
        cursor.execute("SELECT * FROM recipes WHERE type = 'main_meal'")
        recipes = [{'id': r[0], 'name': r[1], 'type': r[2], 'ingredients': r[3].split(','), 'instructions': r[4], 'remarks': r[5], 'image': r[6]} for r in cursor.fetchall()]
        app.logger.debug(f"Main meals fetched: {recipes}")
    return render_template('main_meals.html', recipes=recipes)

@app.route('/baking')
def baking():
    DATABASE_URL = os.environ['DATABASE_URL']
    with psycopg2.connect(DATABASE_URL, sslmode='require') as conn:
        cursor = conn.cursor()
        cursor.execute("SELECT * FROM recipes WHERE type = 'baking'")
        recipes = [{'id': r[0], 'name': r[1], 'type': r[2], 'ingredients': r[3].split(','), 'instructions': r[4], 'remarks': r[5], 'image': r[6]} for r in cursor.fetchall()]
        app.logger.debug(f"Baking recipes fetched: {recipes}")
    return render_template('baking.html', recipes=recipes)

@app.route('/desert')
def desert():
    DATABASE_URL = os.environ['DATABASE_URL']
    with psycopg2.connect(DATABASE_URL, sslmode='require') as conn:
        cursor = conn.cursor()
        cursor.execute("SELECT * FROM recipes WHERE type = 'desert'")
        recipes = [{'id': r[0], 'name': r[1], 'type': r[2], 'ingredients': r[3].split(','), 'instructions': r[4], 'remarks': r[5], 'image': r[6]} for r in cursor.fetchall()]
        app.logger.debug(f"Desert recipes fetched: {recipes}")
    return render_template('desert.html', recipes=recipes)

@app.route('/add_recipe')
def add_recipe_page():
    return render_template('add_recipe.html')

@app.route('/edit_recipe/<int:id>')
def edit_recipe(id):
    DATABASE_URL = os.environ['DATABASE_URL']
    with psycopg2.connect(DATABASE_URL, sslmode='require') as conn:
        cursor = conn.cursor()
        cursor.execute("SELECT * FROM recipes WHERE id = %s", (id,))
        recipe = cursor.fetchone()
        if recipe:
            recipe_data = {
                'id': recipe[0],
                'name': recipe[1],
                'type': recipe[2],
                'ingredients': recipe[3].split(','),
                'instructions': recipe[4],
                'remarks': recipe[5] or '',
                'image': recipe[6] or ''
            }
            app.logger.debug(f"Editing recipe: {recipe_data}")
            return render_template('edit_recipe.html', recipe=recipe_data)
        return jsonify({'status': 'error', 'message': 'Recipe not found'}), 404

@app.route('/edit_recipe/<int:id>', methods=['POST'])
def update_recipe(id):
    try:
        data = request.get_json()
        name = data['name']
        type = data['type']
        ingredients = ','.join(data['ingredients'])
        instructions = data['instructions']
        remarks = data['remarks']
        image = data['image']
        DATABASE_URL = os.environ['DATABASE_URL']
        with psycopg2.connect(DATABASE_URL, sslmode='require') as conn:
            cursor = conn.cursor()
            cursor.execute('''
                UPDATE recipes
                SET name = %s, type = %s, ingredients = %s, instructions = %s, remarks = %s, image = %s
                WHERE id = %s
            ''', (name, type, ingredients, instructions, remarks, image, id))
            conn.commit()
            if cursor.rowcount > 0:
                return jsonify({'status': 'success'}), 200
            else:
                return jsonify({'status': 'error', 'message': 'Recipe not found'}), 404
    except Exception as e:
        app.logger.error(f"Error updating recipe: {str(e)}")
        return jsonify({'status': 'error', 'message': str(e)}), 500

@app.route('/contact')
def contact():
    return render_template('contact.html')

@app.route('/add_recipe', methods=['POST'])
def add_recipe():
    try:
        data = request.get_json()
        name = data['name']
        type = data['type']
        ingredients = ','.join(data['ingredients'])
        instructions = data['instructions']
        remarks = data['remarks']
        image = data['image']
        DATABASE_URL = os.environ['DATABASE_URL']
        with psycopg2.connect(DATABASE_URL, sslmode='require') as conn:
            cursor = conn.cursor()
            cursor.execute('''
                INSERT INTO recipes (name, type, ingredients, instructions, remarks, image)
                VALUES (%s, %s, %s, %s, %s, %s)
            ''', (name, type, ingredients, instructions, remarks, image))
            conn.commit()
        app.logger.debug(f"Added recipe: {name}")
        return jsonify({'status': 'success'}), 200
    except Exception as e:
        app.logger.error(f"Error adding recipe: {str(e)}")
        return jsonify({'status': 'error', 'message': str(e)}), 500

@app.route('/delete_recipe/<int:id>', methods=['DELETE'])
def delete_recipe(id):
    try:
        DATABASE_URL = os.environ['DATABASE_URL']
        with psycopg2.connect(DATABASE_URL, sslmode='require') as conn:
            cursor = conn.cursor()
            cursor.execute('DELETE FROM recipes WHERE id = %s', (id,))
            conn.commit()
            if cursor.rowcount > 0:
                return jsonify({'status': 'success'}), 200
            else:
                return jsonify({'status': 'error', 'message': 'Recipe not found'}), 404
    except Exception as e:
        app.logger.error(f"Error deleting recipe: {str(e)}")
        return jsonify({'status': 'error', 'message': str(e)}), 500

@app.route('/download_recipe/<int:id>/<format>')
def download_recipe(id, format):
    DATABASE_URL = os.environ['DATABASE_URL']
    with psycopg2.connect(DATABASE_URL, sslmode='require') as conn:
        cursor = conn.cursor()
        cursor.execute("SELECT * FROM recipes WHERE id = %s", (id,))
        recipe = cursor.fetchone()
        if not recipe:
            return jsonify({'status': 'error', 'message': 'Recipe not found'}), 404
        recipe_data = {
            'name': recipe[1],
            'type': recipe[2],
            'ingredients': recipe[3].split(','),
            'instructions': recipe[4].split('.'),
            'remarks': recipe[5],
            'image': recipe[6]
        }
    if format == 'pdf':
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter)
        styles = getSampleStyleSheet()
        elements = []
        elements.append(Paragraph(recipe_data['name'], styles['Title']))
        elements.append(Spacer(1, 12))
        elements.append(Paragraph(f"Type: {recipe_data['type'].replace('_', ' ').title()}", styles['Normal']))
        elements.append(Spacer(1, 12))
        elements.append(Paragraph("Ingredients:", styles['Heading2']))
        ingredients_list = ListFlowable(
            [ListItem(Paragraph(ing.strip(), styles['Normal'])) for ing in recipe_data['ingredients']],
            bulletType='1'
        )
        elements.append(ingredients_list)
        elements.append(Spacer(1, 12))
        elements.append(Paragraph("Instructions:", styles['Heading2']))
        instructions_list = ListFlowable(
            [ListItem(Paragraph(ins.strip(), styles['Normal'])) for ins in recipe_data['instructions'] if ins.strip()],
            bulletType='1'
        )
        elements.append(instructions_list)
        elements.append(Spacer(1, 12))
        if recipe_data['remarks']:
            elements.append(Paragraph("Remarks:", styles['Heading2']))
            elements.append(Paragraph(recipe_data['remarks'], styles['Normal']))
        doc.build(elements)
        buffer.seek(0)
        return send_file(buffer, as_attachment=True, download_name=f"{recipe_data['name']}.pdf", mimetype='application/pdf')
    elif format == 'docx':
        doc = Document()
        doc.add_heading(recipe_data['name'], 0)
        doc.add_paragraph(f"Type: {recipe_data['type'].replace('_', ' ').title()}")
        doc.add_heading("Ingredients:", level=2)
        for ing in recipe_data['ingredients']:
            doc.add_paragraph(ing.strip(), style='List Number')
        doc.add_heading("Instructions:", level=2)
        for ins in recipe_data['instructions']:
            if ins.strip():
                doc.add_paragraph(ins.strip(), style='List Number')
        if recipe_data['remarks']:
            doc.add_heading("Remarks:", level=2)
            doc.add_paragraph(recipe_data['remarks'])
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return send_file(buffer, as_attachment=True, download_name=f"{recipe_data['name']}.docx", mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    return jsonify({'status': 'error', 'message': 'Invalid format'}), 400

@app.route('/submit_form', methods=['POST'])
def submit_form():
    try:
        name = request.form.get('name')
        email = request.form.get('email')
        message = request.form.get('message')
        telegram_message = (
            f"Culinary Chaos Message 🌭🍟🍔🍪🎂!\n\n"
            f"From 📛: {name}\n"
            f"Email 📧: {email}\n"
            f"Message ℹ️:{message}"
        )
        bot.send_message(CHAT_ID, telegram_message)
        return """
        <script>
            alert('Message sent successfully!');
            window.location.href = '/';
        </script>
        """
    except Exception as e:
        logger.error(f"Error sending Telegram message: {str(e)}")
        return """
        <script>
            alert('Error sending message. Please try again later.');
            window.location.href = '/';
        </script>
        """

# Serve static files
@app.route('/static/<path:path>')
def send_static(path):
    return send_from_directory('static', path)

if __name__ == '__main__':
    port = int(os.environ.get('PORT', 5100))
    logger.debug(f"Running on port {port}")
    socketio.run(app, host='0.0.0.0', port=port, debug=True)